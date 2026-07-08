from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from typing import Any

from src.models import AggregatedReport, ReportNarrative


def _safe_get_category_val(mapping: dict, category_id: int, default: Any = 0) -> Any:
    if mapping is None:
        return default
    return mapping.get(category_id, mapping.get(str(category_id), default))


TITLE_LINE_1_INDEX = 3
TITLE_LINE_2_INDEX = 4
TITLE_LINE_3_INDEX = 5
PERIOD_LINE_INDEX = 6
NOTE_TITLE_INDEX = 9
SUMMARY_INDEX = 10
CATEGORY_2_INDEX = 12
CATEGORY_3_INDEX = 14
CATEGORY_4_INDEX = 15
CATEGORY_5_INDEX = 16
CATEGORY_6_INDEX = 17
CATEGORY_7_INDEX = 18
CATEGORY_8_INDEX = 19
CATEGORY_9_INDEX = 20
CATEGORY_10_INDEX = 21
CATEGORY_11_INDEX = 22
CATEGORY_12_INDEX = 23
FINAL_HEADING_INDEX = 24
FINAL_OBSERVATION_INDEX = 25


def generate_report_docx(
    reference_report_path: Path,
    output_dir: Path,
    aggregated: AggregatedReport,
    narrative: ReportNarrative,
) -> Path:
    document = Document(reference_report_path)
    _apply_page_setup(document)
    _fill_header_block(document, aggregated)
    _fill_table(document.tables[0], aggregated)
    _fill_narrative(document, narrative)

    output_path = output_dir / _build_report_filename(aggregated)
    document.save(output_path)
    return output_path


def _apply_page_setup(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE


def _fill_header_block(document: Document, aggregated: AggregatedReport) -> None:
    doc_number = "01-03/80"
    report_date = aggregated.period.end.strftime("%d.%m.%Y йил")
    period_label = _format_period_label(aggregated)

    _set_paragraph_text(document.paragraphs[0], doc_number, bold=True, font_size=12)
    _set_paragraph_text(document.paragraphs[1], report_date, bold=True, font_size=12)
    _set_paragraph_text(
        document.paragraphs[TITLE_LINE_1_INDEX],
        "Халқ депутатлари Наманган вилояти, туман ва шаҳар Кенгашларининг расмий телеграм каналларида",
        bold=True,
        font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_paragraph_text(
        document.paragraphs[TITLE_LINE_2_INDEX],
        "эълон қилинган материаллар бўйича",
        bold=True,
        font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_paragraph_text(
        document.paragraphs[TITLE_LINE_3_INDEX],
        "ҲАФТАЛИК МАЪЛУМОТ",
        bold=True,
        font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_paragraph_text(
        document.paragraphs[PERIOD_LINE_INDEX],
        period_label,
        italic=True,
        font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_paragraph_text(document.paragraphs[NOTE_TITLE_INDEX], "Изоҳ:", bold=True, font_size=12)


def _fill_table(table, aggregated: AggregatedReport) -> None:
    for row_index, stats in enumerate(aggregated.channel_stats, start=2):
        row = table.rows[row_index]
        values = [
            str(row_index - 1),
            stats.council_name,
            str(stats.total_posts),
            str(_safe_get_category_val(stats.category_counts, 2)),
            str(_safe_get_category_val(stats.category_counts, 3)),
            str(_safe_get_category_val(stats.category_counts, 4)),
            str(_safe_get_category_val(stats.category_counts, 5)),
            str(_safe_get_category_val(stats.category_counts, 6)),
            str(_safe_get_category_val(stats.category_counts, 7)),
            str(_safe_get_category_val(stats.category_counts, 8)),
            str(_safe_get_category_val(stats.category_counts, 9)),
            str(_safe_get_category_val(stats.category_counts, 10)),
            str(_safe_get_category_val(stats.category_counts, 11)),
            str(_safe_get_category_val(stats.category_counts, 12)),
            str(stats.subscriber_count or 0),
            str(stats.average_views),
        ]
        _fill_row(row, values)

    totals_row = table.rows[len(aggregated.channel_stats) + 2]
    totals = aggregated.totals_by_category
    total_values = [
        "Жами",
        "Жами",
        str(_safe_get_category_val(totals, 1)),
        str(_safe_get_category_val(totals, 2)),
        str(_safe_get_category_val(totals, 3)),
        str(_safe_get_category_val(totals, 4)),
        str(_safe_get_category_val(totals, 5)),
        str(_safe_get_category_val(totals, 6)),
        str(_safe_get_category_val(totals, 7)),
        str(_safe_get_category_val(totals, 8)),
        str(_safe_get_category_val(totals, 9)),
        str(_safe_get_category_val(totals, 10)),
        str(_safe_get_category_val(totals, 11)),
        str(_safe_get_category_val(totals, 12)),
        str(sum((item.subscriber_count or 0) for item in aggregated.channel_stats)),
        str(aggregated.overall_average_views),
    ]
    _fill_row(totals_row, total_values)


def _fill_narrative(document: Document, narrative: ReportNarrative) -> None:
    _set_paragraph_text(document.paragraphs[SUMMARY_INDEX], narrative.summary, font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_2_INDEX], _safe_get_category_val(narrative.category_sections, 2, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_3_INDEX], _safe_get_category_val(narrative.category_sections, 3, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_4_INDEX], _safe_get_category_val(narrative.category_sections, 4, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_5_INDEX], _safe_get_category_val(narrative.category_sections, 5, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_6_INDEX], _safe_get_category_val(narrative.category_sections, 6, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_7_INDEX], _safe_get_category_val(narrative.category_sections, 7, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_8_INDEX], _safe_get_category_val(narrative.category_sections, 8, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_9_INDEX], _safe_get_category_val(narrative.category_sections, 9, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_10_INDEX], _safe_get_category_val(narrative.category_sections, 10, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_11_INDEX], _safe_get_category_val(narrative.category_sections, 11, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[CATEGORY_12_INDEX], _safe_get_category_val(narrative.category_sections, 12, ""), font_size=12)
    _set_paragraph_text(document.paragraphs[FINAL_HEADING_INDEX], narrative.final_heading, font_size=12)
    _set_paragraph_text(document.paragraphs[FINAL_OBSERVATION_INDEX], narrative.final_observation, font_size=12)


def _fill_row(row, values: list[str]) -> None:
    for cell, value in zip(row.cells, values, strict=True):
        _replace_cell_text(cell, value)


def _replace_cell_text(cell, value: str) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    first_run = paragraph.runs[0] if paragraph.runs else None

    for extra_paragraph in list(cell.paragraphs[1:]):
        extra_paragraph._element.getparent().remove(extra_paragraph._element)

    _clear_paragraph(paragraph)
    run = paragraph.add_run(value)
    if first_run is not None:
        if first_run.font.size:
            run.font.size = first_run.font.size
        run.bold = first_run.bold
        run.italic = first_run.italic


def _set_paragraph_text(paragraph, text: str, font_size: int, bold: bool = False, italic: bool = False, alignment=None) -> None:
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if alignment is not None:
        paragraph.alignment = alignment


def _clear_paragraph(paragraph) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag.endswith('}pPr'):
            continue
        paragraph_element.remove(child)


def _format_period_label(aggregated: AggregatedReport) -> str:
    start = aggregated.period.start
    end = aggregated.period.end
    if start.month == end.month and start.year == end.year:
        return f"({start.day} — {end.strftime('%d.%m.%Y')} йил)"
    return f"({start.strftime('%d.%m.%Y')} — {end.strftime('%d.%m.%Y')} йил)"


def _build_report_filename(aggregated: AggregatedReport) -> str:
    start_label = aggregated.period.start.strftime("%Y-%m-%d")
    end_label = aggregated.period.end.strftime("%Y-%m-%d")
    return f"telegram_monitoring_{start_label}_{end_label}.docx"
