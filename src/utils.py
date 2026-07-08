from __future__ import annotations

import json
import logging
from pathlib import Path

from docx import Document

from src.models import RunDirectories


def configure_logging() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    )


def ensure_run_directories(project_root: Path, output_root: Path) -> RunDirectories:
    root_dir = project_root / output_root
    raw_dir = root_dir / "raw_posts"
    classified_dir = root_dir / "classified_posts"
    regression_dir = root_dir / "regression_checks"
    audited_dir = root_dir / "audited_posts"
    aggregated_dir = root_dir / "aggregated_stats"
    narrative_dir = root_dir / "narrative"
    reports_dir = root_dir / "reports"

    for path in [raw_dir, classified_dir, regression_dir, audited_dir, aggregated_dir, narrative_dir, reports_dir]:
        path.mkdir(parents=True, exist_ok=True)

    return RunDirectories(
        root_dir=root_dir,
        raw_posts_path=raw_dir / "raw_posts.json",
        classified_posts_path=classified_dir / "classified_posts.json",
        classification_regression_checks_path=regression_dir / "classification_regression_checks.json",
        audited_posts_path=audited_dir / "audited_posts.json",
        audit_regression_checks_path=regression_dir / "audit_regression_checks.json",
        aggregated_stats_path=aggregated_dir / "aggregated_stats.json",
        narrative_path=narrative_dir / "narrative.json",
        narrative_validation_path=narrative_dir / "narrative_validation.json",
        reports_dir=reports_dir,
    )


def write_json(path: Path, payload: object) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_utf8_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def chunked[T](items: list[T], size: int) -> list[list[T]]:
    return [items[index : index + size] for index in range(0, len(items), size)]
