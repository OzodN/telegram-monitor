from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def summarize_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.replace("\n", " ").strip()
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name,
                "text": text,
            }
        )

    tables = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows, start=1):
            rows.append(
                {
                    "index": row_index,
                    "cells": [cell.text.replace("\n", " ").strip() for cell in row.cells],
                }
            )
        tables.append(
            {
                "index": table_index,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "rows": rows,
            }
        )

    return {
        "name": path.name,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    payload = [summarize_docx(Path(arg)) for arg in sys.argv[1:]]
    json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
