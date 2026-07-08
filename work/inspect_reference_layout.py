from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def twips_to_inches(value: int | None) -> float | None:
    if value is None:
        return None
    return round(value / 1440, 3)


def emu_to_inches(value: int | None) -> float | None:
    if value is None:
        return None
    return round(value / 914400, 3)


def paragraph_runs(paragraph) -> list[dict]:
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue
        font = run.font
        runs.append(
            {
                "text": run.text,
                "bold": bool(font.bold) if font.bold is not None else None,
                "italic": bool(font.italic) if font.italic is not None else None,
                "underline": bool(font.underline) if font.underline is not None else None,
                "font_name": font.name,
                "font_size_pt": round(font.size.pt, 1) if font.size else None,
            }
        )
    return runs


def main() -> None:
    path = Path(r"D:\ozod\coding\python\telegram-automation-system\prompt\Телеграм_бўйича_ҳафталик_жадвал_22_28_июнь.docx")
    doc = Document(path)

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width_in": emu_to_inches(section.page_width),
                "page_height_in": emu_to_inches(section.page_height),
                "left_margin_in": emu_to_inches(section.left_margin),
                "right_margin_in": emu_to_inches(section.right_margin),
                "top_margin_in": emu_to_inches(section.top_margin),
                "bottom_margin_in": emu_to_inches(section.bottom_margin),
                "header_distance_in": emu_to_inches(section.header_distance),
                "footer_distance_in": emu_to_inches(section.footer_distance),
            }
        )

    interesting_paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.replace("\n", " ").strip()
        if not text:
            continue
        fmt = paragraph.paragraph_format
        interesting_paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name,
                "alignment": str(paragraph.alignment),
                "space_before_pt": round(fmt.space_before.pt, 1) if fmt.space_before else None,
                "space_after_pt": round(fmt.space_after.pt, 1) if fmt.space_after else None,
                "line_spacing": fmt.line_spacing,
                "left_indent_in": twips_to_inches(fmt.left_indent.twips) if fmt.left_indent else None,
                "first_line_indent_in": twips_to_inches(fmt.first_line_indent.twips) if fmt.first_line_indent else None,
                "text": text,
                "runs": paragraph_runs(paragraph),
            }
        )

    table = doc.tables[0]
    table_summary = {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "autofit": table.autofit,
        "first_row_cell_widths_in": [emu_to_inches(cell.width) for cell in table.rows[0].cells],
    }

    payload = {
        "sections": sections,
        "paragraphs": interesting_paragraphs,
        "table": table_summary,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
